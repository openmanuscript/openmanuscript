from . import core

import os
import json
import re
# import commonmark
import markdown as MDown
import datetime
import time

from .html import add_html

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn

# -----------------------------------------------------------------------------
# globals
# -----------------------------------------------------------------------------

MARGIN = {
    "page" : 1,
    "toc"  : 1.5
}

SETTINGS = {
    "paragraph_before" : 0,
    "paragraph_after"  : 0,
    "mark_parser"      : "python-markdown"
}

PART_STATE = {
    "bold"   : False,
    "italic" : False
}

# the following three fuctions add page numbers, thanks to:
# https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
#
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

# set the starting page number of a section
def set_start_page_number(section, number=1, fmt='decimal'):
    pagenum  = create_element('w:pgNumType')
    create_attribute(pagenum, 'w:start', str(number).encode("utf-8"))
    create_attribute(pagenum, 'w:fmt', fmt) 
    section._sectPr.append(pagenum)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
#
# End of stackoverflow solution. Thanks!
#

def write_toc( doc, ms ):
    # doc.add_page_break()
    add_toc_section(doc, WD_SECTION.CONTINUOUS)

    # spacing from top of page 
    doc.add_paragraph()

    # table of contents text 
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.left_indent = Pt(0)
    run = p.add_run("Table of Contents")

    # spacing to chapters
    for i in range(2):
        doc.add_paragraph()

    # entries for chapters
    cur_chapter = 1
    for chapter in ms["chapters"]:
        chaptype = core.get_chapter_type(chapter)
        if core.check_chapter_tags(chapter) and (chaptype != "QUOTE"):
            chapnum = "CHAPTER {}".format(cur_chapter).upper()
            chaptername = ""
            increment_chapter = True
            if "title" in chapter:
                chaptername = chapter["title"]

            if (chaptype == "CHAPTER"):
                allcaps_title = chapnum
            else:
                if (chaptype == "PROLOGUE"):
                    allcaps_title = chaptype
                else:
                    allcaps_title = ""
                increment_chapter = False

            if increment_chapter:
                cur_chapter += 1

            p = doc.add_paragraph()
            pf = p.paragraph_format
            if (chaptype == "PART"):
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf.space_before = Inches(0.25)
            else:
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.space_before = Inches(0)
                pf.space_after = Inches(0)

            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            pf.left_indent = Inches(0.25)
            pf.first_line_indent = Inches(-0.25)
            add_tab_stop( p, Inches(1.25) )
            if allcaps_title == "":
                run = p.add_run(chaptername)
            else:
                run = p.add_run("{}: \t{}".format(allcaps_title, chaptername))

    # end the TOC
    doc.add_page_break()
    

def add_to_run(run, text):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = text 

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_page_slug_header(section):
    # section.different_first_page_header = True
    header = section.header
    p = header.add_paragraph()
    p.add_run("{}/ {}/ ".format(core.author["surname"], core.
                manuscript["runningtitle"].upper()))
    run = p.add_run()
    add_page_number(run)

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def debug_sections(document):
    print("WORD: debugging sections ...")
    cur = 0
    for section in document.sections:
        print("Section {}".format(cur))
        print("{}".format(section.start_type))
        cur = cur + 1

def add_toc_section(document, s_type):
    new_section = document.add_section(s_type)

    # set margins 
    new_section.top_margin    = Inches(MARGIN["page"])
    new_section.bottom_margin = Inches(MARGIN["page"])
    new_section.left_margin   = Inches(MARGIN["toc"])
    new_section.right_margin  = Inches(MARGIN["toc"])
    header = new_section.header
    header.is_linked_to_previous = False

def add_main_section(document, s_type):
    new_section = document.add_section(s_type)

    # set margins 
    new_section.top_margin    = Inches(MARGIN["page"])
    new_section.bottom_margin = Inches(MARGIN["page"])
    new_section.left_margin   = Inches(MARGIN["page"])
    new_section.right_margin  = Inches(MARGIN["page"])
    header = new_section.header
    header.is_linked_to_previous = False

    set_start_page_number(new_section) 
    add_page_slug_header(new_section)

def write_title(document):
    # word count
    count = core.get_approximate_word_count() 
    p = document.add_paragraph("Approx. words: {:,}".format(count))
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # author information
    p = document.add_paragraph("{}\n{}\n{}, {} {}\n{}\n{}\n".format(
                            core.author["name"],
                            core.author["streetAddress"],
                            core.author["addressLocality"], 
                            core.author["addressRegion"], 
                            core.author["postalCode"],
                            core.author["phone"],
                            core.author["email"]))
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # spacing to title
    for i in range(7):
        document.add_paragraph()

    # title
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(core.manuscript["title"])

    # by 
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run("by")

    # author 
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run("{}".format(core.author["name"]))

    # slug 
    if (core.settings["slug"] != None) and (core.settings["slug"] != "None"):
        p = document.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        run = p.add_run("({})".format(core.settings["slug"]))

    # end the page
    document.add_page_break()


def write_paragraph_space(document):
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

def write_chapter_heading(document, chapter, chapnum, chaptype):
    chapnum = "CHAPTER {0}".format(chapnum).upper()
    increment_chapter = True

    chaptername = ""
    if "title" in chapter:
        chaptername = chapter["title"]

    if (chaptype == "CHAPTER"):
        allcaps_title = chapnum
    else:
        if (chaptype == "PROLOGUE"):
            allcaps_title = chaptype 
        else:
            allcaps_title = "" 
        increment_chapter = False

    # add half a page of spacing
    # document.add_page_break()
    for i in range(0, 10):
        p = document.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # add title
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(allcaps_title)
    run.bold = True

    # add chapter name
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(chaptername)
    run.bold = True

    p = document.add_paragraph()

    return increment_chapter



def write_preamble(doc):
    # Set up the document
    style     = doc.styles['Normal']
    font      = style.font
    font.name = core.settings["font"]
    font.size = Pt(int(core.settings["fontsize"]))

    # Styles
    # Body
    styleIndent = doc.styles.add_style('Body', WD_STYLE_TYPE.PARAGRAPH)
    font      = styleIndent.font
    font.name = core.settings["font"]
    font.size = Pt(int(core.settings["fontsize"]))
    pf = styleIndent.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(0.5)
    # pf.space_before = Pt(int(core.settings["fontsize"]))
    pf.space_before  = SETTINGS["paragraph_before"] 
    pf.space_after   = SETTINGS["paragraph_after"] 
    pf.widow_control = True

    # Heading1
    for i in range(1, 10):
        curstyle  = doc.styles.add_style('Heading{}'.format(i), WD_STYLE_TYPE.PARAGRAPH)
        font      = curstyle.font
        font.name = core.settings["font"]
        font.size = Pt(int(core.settings["fontsize"]))
        font.bold = True
        pf = curstyle.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Inches(0.0)
        pf.space_before  = SETTINGS["paragraph_before"] 
        pf.space_after   = SETTINGS["paragraph_after"] 
        pf.widow_control = True

    # List items
    style     = doc.styles['List Bullet']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style     = doc.styles['List Number']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

def write_postamble(doc):
    # spacing to title
    for i in range(15):
        doc.add_paragraph()

    # add text
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run("{}".format(core.postamble))

    return

def write_docinfo(doc):
    p = doc.core_properties
    p.author   = core.author["name"]
    p.created  = datetime.datetime.now()
    p.title    = core.manuscript["title"]
    p.comments = "created by {} v{}".format(core.get_name(), core.get_version())
    return

# -----------------------------------------------------------------------------
# remove comments 
# -----------------------------------------------------------------------------
def remove_comments( data ):
    return handle_tag( data, "comment", False) 

# -----------------------------------------------------------------------------
# remove notes 
# -----------------------------------------------------------------------------
def handle_notes( data, state ):
    return handle_tag( data, "notes", state)

# -----------------------------------------------------------------------------
# include a file 
# -----------------------------------------------------------------------------
def handle_include( data ):
    includes = re.findall(r'(<include>\s*([^<]+).md\s*</include>)', data, flags=re.DOTALL)

    for i in includes:
        scenefile = core.get_scenefile(i[1])
        with open(scenefile, "r") as s_file:
            scenetext = s_file.read()
            isub = re.compile(i[0], re.DOTALL)
            data = re.sub(isub, scenetext, data)

    return data

# -----------------------------------------------------------------------------
# substitute emdash 
# -----------------------------------------------------------------------------
def handle_emdash( data ):
    subflags = flags=re.MULTILINE|re.DOTALL
    data  = re.sub(rf'--', u'\u2014', data, flags=subflags)

    return data


# -----------------------------------------------------------------------------
# handle arbitrary html syntax tag 
# -----------------------------------------------------------------------------
def handle_tag( data, tag, include ):
    subflags = flags=re.MULTILINE|re.DOTALL
    if include:
        # include the text surrounded by the tag, but remote the tag
        data  = re.sub(rf'<{tag}>', "", data, flags=subflags)
        data  = re.sub(rf'</{tag}>', "", data, flags=subflags) 
    else:
        # remove the text surrounded by the tag, as well as the tag
        data = re.sub(rf'<{tag}>.*?</{tag}>', "", data, flags=subflags) 

    return data

# -----------------------------------------------------------------------------
# write a scene separator
# -----------------------------------------------------------------------------
def write_scene_separator(doc, scene):
    separator = "###"

    if core.settings["filescenesep"]:
        separator = scene

    for i in range(1):
        doc.add_paragraph()

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(separator)

    for i in range(1):
        doc.add_paragraph()

# -----------------------------------------------------------------------------
# write a single scene
# -----------------------------------------------------------------------------
def write_scenefile(doc, scene):
    scenefile = core.get_scenefile(scene)
    if os.path.isfile(scenefile):
        pgraph = doc.add_paragraph()
        pgraph.style = 'Body'

        # read, transform and add the scene file 
        with open(scenefile, "r") as s_file:
            scenetext = s_file.read()
            html_tree = create_scene_text(scenetext)

            if (html_tree != None):
                # print(html_tree)
                add_html(pgraph, html_tree)

    else:
        print("ERROR: can't find scene file: " + scenefile)
        return 0

# -----------------------------------------------------------------------------
# write a single scene
# -----------------------------------------------------------------------------
def write_scenetext(doc, scenetext):
    # print("in write_scenetext")
    pgraph = doc.add_paragraph()
    pgraph.style = 'Body'

    html_tree = create_scene_text(scenetext)

    if (html_tree != None):
        # print(html_tree)
        add_html(pgraph, html_tree)

    else:
        print("ERROR: can't add scene text: " + scenetext)
        return 0

# -----------------------------------------------------------------------------
# transform raw input text into text that can be added to a paragraph 
#   this includes both cleaning up whitespace and transforming into html
# -----------------------------------------------------------------------------
def create_scene_text(scenetext):
    scenetext = scenetext.strip()
    scenetext = remove_comments(scenetext)

    if (core.settings["excludesections"] != None):
        for section in core.settings["excludesections"]:
            scenetext = handle_tag(scenetext, section, False) 

    if (core.settings["includesections"] != None):
        for section in core.settings["includesections"]:
            scenetext = handle_tag(scenetext, section, True) 

    scenetext = handle_include(scenetext)
    scenetext = handle_notes(scenetext, core.settings["notes"])
    scenetext = handle_emdash(scenetext)

    html_tree = None
    if scenetext:
        if (SETTINGS["mark_parser"] == "commonmark"):
            # commonmark
                # does not handle footnotes
            html_tree = commonmark.commonmark(scenetext)
        elif (SETTINGS["mark_parser"] == "python-markdown"):
            # python-markdown
                # handles footnotes
            html_tree = MDown.markdown(scenetext, extensions=['footnotes'])
        else:
            print("ERROR: unsupported markdown parser")
            exit(1)

    return html_tree

# -----------------------------------------------------------------------------
# write a single chapter 
# -----------------------------------------------------------------------------
def write_chapter(doc, chapter, chapnum):
    increment_chapter = False

    chaptype = core.get_chapter_type(chapter)

    scenes = None
    if core.settings["manuscripttype"] == "novel":
        increment_chapter = write_chapter_heading(doc, chapter, chapnum, chaptype)
    elif core.settings["manuscripttype"] == "story":
        write_paragraph_space(doc)

    # write content
    first = True
    scenes = [] 
    if core.settings["chaptersummary"]:
        if "summary" in chapter:
            scenes = [chapter["summary"]]
    else:
        if "scenes" in chapter:
            scenes = chapter["scenes"]

    if core.settings["chapterdesc"]:
        # just write the description of the chapter
        if "desc" in chapter:
            write_scenetext(doc, chapter["desc"])
    else:
        # write out all the scenes
        for scene in scenes: 
            if not first:
                write_scene_separator(doc, scene)
            else:
                if core.settings["filescenesep"]:
                    write_scene_separator(doc, scene)
                first = False

            if not scene.endswith(".pdf"):
                write_scenefile(doc, scene)

    # end the chapter
    doc.add_page_break()

    return increment_chapter

def write_chapters(doc, manuscript):
    add_main_section(doc, WD_SECTION.CONTINUOUS)

    # now add the sections
    chapnum = 1

    incr_chapter = True
    for chapter in manuscript["chapters"]:
        if core.check_chapter_tags(chapter):
            incr_chapter = write_chapter(doc, chapter, chapnum)

            if (incr_chapter):
                chapnum += 1

def add_tab_stop( p, location ):
    tab_stops = p.paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(location, WD_TAB_ALIGNMENT.LEFT)

def write(outputfile):

    # create the renderer, and get started
    result = False 
    with open( outputfile, "w") as f:
        result = True

        # create the one document
        theDocument  = Document()

        # construct the document
        write_preamble(theDocument)
        write_docinfo(theDocument)
        # write_headers(theDocument)
        if not core.settings["notitlepage"]:
            write_title(theDocument)
        if core.settings["toc"]:
            write_toc(theDocument, core.manuscript)
        # debug_sections(theDocument)

        write_chapters(theDocument, core.manuscript)
        write_postamble(theDocument)

        theDocument.save(outputfile)

    return result



