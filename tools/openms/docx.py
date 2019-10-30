from . import core

import os
import json
import re
# import commonmark
import markdown as MDown
import datetime
import time

from .html import add_html

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn

# -----------------------------------------------------------------------------
# globals
# -----------------------------------------------------------------------------

MARGIN = {
    "page" : 1
}

SETTINGS = {
    "paragraph_before" : 0,
    "paragraph_after"  : 0,
    "mark_parser"      : "python-markdown"
}

PART_STATE = {
    "bold"   : False,
    "italic" : False
}

# the following three fuctions add page numbers, thanks to:
# https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
#
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
#
# End of stackoverflow solution. Thanks!
#

def add_to_run(run, text):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = text 

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_page_slug_header(section):
    # section.different_first_page_header = True
    header = section.header
    p = header.add_paragraph()
    p.add_run("{}/ {}/ ".format(core.author["surname"], core.
                manuscript["runningtitle"].upper()))
    run = p.add_run()
    add_page_number(run)

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def debug_sections(document):
    print("WORD: debugging sections ...")
    cur = 0
    for section in document.sections:
        print("Section {}".format(cur))
        print("{}".format(section.start_type))
        cur = cur + 1

def add_section(document, s_type):
    new_section = document.add_section(s_type)
    # set margins 
    new_section.top_margin    = Inches(MARGIN["page"])
    new_section.bottom_margin = Inches(MARGIN["page"])
    new_section.left_margin   = Inches(MARGIN["page"])
    new_section.right_margin  = Inches(MARGIN["page"])
    header = new_section.header
    header.is_linked_to_previous = False

    add_page_slug_header(new_section)

def write_title(document):
    p = document.add_paragraph("{}\n{}\n{}, {} {}\n{}\n{}\n".format(
                            core.author["name"],
                            core.author["streetAddress"],
                            core.author["addressLocality"], 
                            core.author["addressRegion"], 
                            core.author["postalCode"],
                            core.author["phone"],
                            core.author["email"]))

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for i in range(7):
        document.add_paragraph()

    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(core.manuscript["title"])

    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run("by")

    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run("{}".format(core.author["name"]))

    # set page margins
    # for now, this should be the only section
    for section in document.sections:
        section.top_margin    = Inches(MARGIN["page"])
        section.bottom_margin = Inches(MARGIN["page"])
        section.left_margin   = Inches(MARGIN["page"])
        section.right_margin  = Inches(MARGIN["page"])

    add_section(document, WD_SECTION.CONTINUOUS)

def write_chapter_heading(document, chapter, chapnum, chaptype):
    chapnum = "CHAPTER {0}".format(chapnum).upper()
    increment_chapter = True

    chaptername = ""
    if "title" in chapter:
        chaptername = chapter["title"]

    if (chaptype == "CHAPTER"):
        allcaps_title = chapnum
    else:
        allcaps_title = "" 
        increment_chapter = False

    # add half a page of spacing
    document.add_page_break()
    for i in range(0, 10):
        p = document.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # add title
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(allcaps_title)
    run.bold = True

    # add chapter name
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(chaptername)
    run.bold = True

    p = document.add_paragraph()

    return increment_chapter


def write_preamble(doc):
    # Set up the document
    style     = doc.styles['Normal']
    font      = style.font
    font.name = core.settings["font"]
    font.size = Pt(int(core.settings["fontsize"]))

    # Styles
    # Body
    styleIndent = doc.styles.add_style('Body', WD_STYLE_TYPE.PARAGRAPH)
    font      = styleIndent.font
    font.name = core.settings["font"]
    font.size = Pt(int(core.settings["fontsize"]))
    pf = styleIndent.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(0.5)
    # pf.space_before = Pt(int(core.settings["fontsize"]))
    pf.space_before  = SETTINGS["paragraph_before"] 
    pf.space_after   = SETTINGS["paragraph_after"] 
    pf.widow_control = True

    # Heading1
    for i in range(1, 10):
        curstyle  = doc.styles.add_style('Heading{}'.format(i), WD_STYLE_TYPE.PARAGRAPH)
        font      = curstyle.font
        font.name = core.settings["font"]
        font.size = Pt(int(core.settings["fontsize"]))
        font.bold = True
        pf = curstyle.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Inches(0.0)
        pf.space_before  = SETTINGS["paragraph_before"] 
        pf.space_after   = SETTINGS["paragraph_after"] 
        pf.widow_control = True

    # List items
    style     = doc.styles['List Bullet']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style     = doc.styles['List Number']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

def write_postamble(doc):
    # print("WORD: write_postamble (no-op)")
    return

def write_docinfo(doc):
    p = doc.core_properties
    p.author   = core.author["name"]
    p.created  = datetime.datetime.now()
    p.title    = core.manuscript["title"]
    p.comments = "created by {} v{}".format(core.get_name(), core.get_version())
    return

# -----------------------------------------------------------------------------
# remove comments 
# -----------------------------------------------------------------------------
def remove_comments( data ):
    data  = re.sub('\<comment\>.*\<\/comment\>', '', data, re.DOTALL)
    return data


# -----------------------------------------------------------------------------
# remove notes 
# -----------------------------------------------------------------------------
def handle_notes( data, state ):
    if state:
        # include the note text
        data  = data.replace("<notes>", "") 
        data  = data.replace("</notes>", "") 
    else:
        # remove the note text
        s = re.compile("<notes>.*?</notes>", re.DOTALL)
        data = re.sub(s, "", data)

    return data

# -----------------------------------------------------------------------------
# write a scene separator
# -----------------------------------------------------------------------------
def write_scene_separator(doc, scene):
    separator = "###"

    if core.settings["filescenesep"]:
        separator = scene

    for i in range(1):
        doc.add_paragraph()

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(separator)

    for i in range(1):
        doc.add_paragraph()

# -----------------------------------------------------------------------------
# write a single scene
# -----------------------------------------------------------------------------
def write_scene(doc, scene):
    scenefile = core.get_scenefile(scene)
    if os.path.isfile(scenefile):
        pgraph = doc.add_paragraph()
        pgraph.style = 'Body'

        # read and clean up the input data - mostly newlines and spaces
        with open(scenefile, "r") as s_file:
            scenetext = s_file.read()
            scenetext = scenetext.strip()
            # scenetext = remove_comments(scenetext)
            scenetext = handle_notes(scenetext, core.settings["notes"])

            if scenetext:
                html_tree = ""
                if (SETTINGS["mark_parser"] == "commonmark"):
                    # commonmark
                        # does not handle footnotes
                    html_tree = commonmark.commonmark(scenetext)
                elif (SETTINGS["mark_parser"] == "python-markdown"):
                    # python-markdown
                        # handles footnotes
                    html_tree = MDown.markdown(scenetext, extensions=['footnotes'])
                else:
                    print("ERROR: unsupported markdown parser")
                    exit(1)

                # print(html_tree)
                add_html(pgraph, html_tree)

    else:
        print("ERROR: can't find scene file: " + scenefile)
        return 0

# -----------------------------------------------------------------------------
# write a single chapter 
# -----------------------------------------------------------------------------
def write_chapter(doc, chapter, chapnum):
    increment_chapter = False

    chaptype = None
    if "type" in chapter:
        chaptype = chapter["type"].upper()
    else:
        chaptype = "CHAPTER"

    scenes = None
    if True:
        increment_chapter = write_chapter_heading(doc, chapter, chapnum, chaptype)

        # write content
        first = True
        for scene in chapter["scenes"]: 
            if not first:
                write_scene_separator(doc, scene)
            else:
                if core.settings["filescenesep"]:
                    write_scene_separator(doc, scene)
                first = False

            if not scene.endswith(".pdf"):
                write_scene(doc, scene)

    return increment_chapter

def write_chapters(doc, manuscript):
    chapnum = 1

    incr_chapter = True
    for chapter in manuscript["chapters"]:
        if core.check_chapter_tags(chapter):
            if core.settings["chaptersummary"]:
                if "summary" in chapter:
                    incr_chapter = write_chapter(doc, chapter["summary"], chapnum) 
                else:
                    # not sure if this is the right thing to do ...
                    print("Chapter summary, but no summary present")
                    # incr_chapter = write_chapter(doc, None, chapnum) 
            else:
                incr_chapter = write_chapter(doc, chapter, chapnum)

            if (incr_chapter):
                chapnum += 1

def write(outputfile):

    # create the renderer, and get started
    with open( outputfile, "w") as f:
        success = True

        # create the one document
        theDocument  = Document()

        # construct the document
        write_preamble(theDocument)
        write_docinfo(theDocument)
        # write_headers(theDocument)
        write_title(theDocument)
        # debug_sections(theDocument)

        write_chapters(theDocument, core.manuscript)
        write_postamble(theDocument)

        theDocument.save(outputfile)

        if (success == True):
            print("wrote file: " + core.settings["outputfile"])



