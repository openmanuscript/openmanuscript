from . import core

import os
import json
import re

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn

from mistletoe import Document as MistletoeDocument
from mistletoe.oms_renderer import OMSRenderer

# -----------------------------------------------------------------------------
# globals
# -----------------------------------------------------------------------------

MARGIN = {
    "page" : 1
}

PART_STATE = {
    "bold"   : False,
    "italic" : False
}

TheRenderer = None

# the following three fuctions add page numbers, thanks to:
# https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
#
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
#
# End of stackoverflow solution. Thanks!
#

def add_to_run(run, text):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = text 

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# The following function adds a hyperlink. Thanks to:
# https://stackoverflow.com/questions/48374357/how-to-add-hyperlink-to-an-image-in-python-docx 
#
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
#
# End of stackoverflow example. Thanks!
#

def add_page_slug_header(section):
    # section.different_first_page_header = True
    header = section.header
    p = header.add_paragraph()
    p.add_run("{}/ {}/ ".format(core.author["surname"], core.
                manuscript["runningtitle"].upper()))
    run = p.add_run()
    add_page_number(run)

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def debug_sections(document):
    print("WORD: debugging sections ...")
    cur = 0
    for section in document.sections:
        print("Section {}".format(cur))
        print("{}".format(section.start_type))
        cur = cur + 1

def add_section(document, s_type):
    new_section = document.add_section(s_type)
    # set margins 
    new_section.top_margin    = Inches(MARGIN["page"])
    new_section.bottom_margin = Inches(MARGIN["page"])
    new_section.left_margin   = Inches(MARGIN["page"])
    new_section.right_margin  = Inches(MARGIN["page"])
    header = new_section.header
    header.is_linked_to_previous = False

    add_page_slug_header(new_section)

def write_title(document):
    p = document.add_paragraph("{}\n{}\n{}, {} {}\n{}\n{}\n".format(
                            core.author["name"],
                            core.author["streetAddress"],
                            core.author["addressLocality"], 
                            core.author["addressRegion"], 
                            core.author["postalCode"],
                            core.author["phone"],
                            core.author["email"]))

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for i in range(0, 9):
        document.add_paragraph()

    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(core.manuscript["title"])

    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run("by")

    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run("{}".format(core.author["name"]))

    # set page margins
    # for now, this should be the only section
    for section in document.sections:
        section.top_margin    = Inches(MARGIN["page"])
        section.bottom_margin = Inches(MARGIN["page"])
        section.left_margin   = Inches(MARGIN["page"])
        section.right_margin  = Inches(MARGIN["page"])

def write_chapter_heading(document, chapter, chapnum, chaptype):
    chapnum = "CHAPTER {0}".format(chapnum).upper()
    document.add_page_break()

    for i in range(0, 10):
        p = document.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if (chaptype == "CHAPTER"):
        allcaps_title = chapnum
        chaptername   = chapter["title"]
    else:
        allcaps_title = chaptype
        chaptername   = ""

    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(allcaps_title)
    run.bold = True

    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(chaptername)
    run.bold = True

    p = document.add_paragraph()


def write_preamble(doc):
    # Set up the document
    style     = doc.styles['Normal']
    font      = style.font
    font.name = core.settings["font"]
    font.size = Pt(int(core.settings["fontsize"]))

    # Styles
    styleIndent = doc.styles.add_style('Indent', WD_STYLE_TYPE.PARAGRAPH)
    font      = styleIndent.font
    font.name = core.settings["font"]
    font.size = Pt(int(core.settings["fontsize"]))
    pf = styleIndent.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(0.5)
    pf.space_before = Pt(12)
    pf.widow_control = True

def write_postamble(doc):
    print("WORD: write_postamble (no-op)")

def write_docinfo(doc):
    print("WORD: write_docinfo (no-op)")

def write_chaptersummary(doc, chapter, chapnum, chaptype):
    print("WORD: write_chaptersummary (no-op)")

# -----------------------------------------------------------------------------
# write a scene separator
# -----------------------------------------------------------------------------
def write_scene_separator(doc, scene):
    separator = "###"

    if core.settings["filescenesep"]:
        separator = scene

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(separator)


# -----------------------------------------------------------------------------
# write a single scene
# -----------------------------------------------------------------------------
def write_scene(doc, scene):
    global TheRenderer

    scenefile = core.get_scenefile(scene)
    if os.path.isfile(scenefile):
        p = doc.add_paragraph()
        TheRenderer.TheParagraph = p
        with open(scenefile, "r") as s_file:
            scenetext = s_file.read()
            scenetext = scenetext.strip()
            rendered = TheRenderer.render(MistletoeDocument(scenetext))
    else:
        print("ERROR: can't find scene file: " + scenefile)
        return 0

# -----------------------------------------------------------------------------
# remove comments 
# -----------------------------------------------------------------------------
def handle_comments( data ):
    data  = re.sub('\[comment\]\:\s*\*\s*\([^\)]*\)', '', data, re.DOTALL)
    return data

def write_chapter(doc, chapter, chapnum, chaptype="CHAPTER"):
    write_chapter_heading(doc, chapter, chapnum, chaptype)

    # write content
    first = True
    for scene in chapter["scenes"]:
        if not first:
            write_scene_separator(doc, scene)
        else:
            if core.settings["filescenesep"]:
                write_scene_separator(doc, scene)
            first = False

        if not scene.endswith(".pdf"):
            write_scene(doc, scene)

def write_chapters(doc, manuscript):
    chapnum = 1

    for chapter in manuscript["chapters"]:
        if core.check_chapter_tags(chapter, core.settings["tags"]):
            if core.is_prologue(chapter) or core.is_epilogue(chapter):
                chaptype = chapter["type"].upper()
                if core.settings["chaptersummary"]:
                    write_chaptersummary(f, chapter, chapnum, chaptype)
                else:
                    write_chapter(doc, chapter, chapnum, chaptype)
            else:
                if core.settings["chaptersummary"]:
                    write_chaptersummary(f, chapter, chapnum)
                else:
                    write_chapter(doc, chapter, chapnum)
                chapnum += 1


def write(outputfile):
    global TheRenderer

    # create the renderer, and get started
    with OMSRenderer() as TheRenderer: 
        with open( outputfile, "w") as f:
            success = True

            # create the one document
            theDocument  = Document()
            TheRenderer.Verbose = False
            TheRenderer.TheDocument = theDocument

            # construct the document
            write_preamble(theDocument)
            write_docinfo(theDocument)
            # write_headers(theDocument)
            write_title(theDocument)
            add_section(theDocument, WD_SECTION.CONTINUOUS)
            # debug_sections(theDocument)
            write_chapters(theDocument, core.manuscript)
            write_postamble(theDocument)

            theDocument.save(outputfile)

            if (success == True):
                print("wrote file: " + core.settings["outputfile"])



