import os
import json
import csv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING


__oms = {
    "toolversion" : "1.0",
    "specversion" : "1.0"
}


settings = {
    "authorfile"     : "author.json",
    "manuscriptfile" : "manuscript.json",
    "manuscriptdir"  : "."
}

author = {
}

manuscript = {
}

def set(attribute, value):
    global settings
    settings[attribute] = value

def set_manuscriptdir( msdir ):
    global settings
    settings["manuscriptdir"] = msdir

def set_manuscriptfile( mfile ):
    global settings
    settings["manuscriptfile"] = mfile 

def set_authorfile( afile ):
    global settings
    settings["authorfile"] = afile 

def get_version():
    global __oms
    return __oms["toolversion"]

def get_spec_version():
    global __oms
    return __oms["specversion"]

def get_authorfile():
    global settings
    return os.path.join( settings["manuscriptdir"], settings["authorfile"] )

def get_manuscriptfile():
    global settings
    return os.path.join( settings["manuscriptdir"], settings["manuscriptfile"] )

def get_author():
    global author
    return author

def get_manuscript():
    global manuscript
    return manuscript

def check_version( json_data ):
    global __oms

    result = 0

    if "version" in json_data:
        if json_data["version"] == __oms["specversion"]:
            result = 1
        else:
            print("ERROR: unsupported openmanuscript version: {}"
                    .format(json_data["version"]))
    else:
        print("ERROR: invalid json data (no version number)") 

    return result


def read_data():
    global author
    global manuscript

    with open( get_authorfile() ) as author_file:
        author = json.load( author_file )
        if (check_version(author)):
            author = author["author"]

    with open( get_manuscriptfile() ) as manuscript_file:
        manuscript = json.load( manuscript_file )
        if (check_version(manuscript)):
            manuscript = manuscript["manuscript"]
        

def csv_to_manuscript( csvfile, ms ):
    with open(csvfile, "r") as csvdata:
        csvreader = csv.reader(csvdata, delimiter=',')

        with open(ms, "w") as mfile:

            mfile.write("{\n")
            mfile.write("\"version\" : \"1.0\",\n")
            mfile.write("\"manuscript\" : {\n")
            mfile.write("    \"title\" : \"Sample\",\n")
            mfile.write("    \"runningtitle\" : \"sample\",\n")
            mfile.write("    \"chapters\" : [\n")

            count = 0
            first = 1
            for row in csvreader:
                if (count > 1):
                    mfile.write(",\n")

                if (count == 0):
                    names = row.copy()
                else:
                    values = row.copy()

                    mfile.write("         {\n")
                    for i in range(len(names)): 
                        mfile.write("             \"{}\" : \"{}\"".format(names[i].strip(), values[i].strip()))
                        if (i == (len(names)-1)):
                            mfile.write("\n")
                        else:
                            mfile.write(",\n")
                if (count != 0):
                    mfile.write("         }")
                count += 1 

            mfile.write("\n")
            mfile.write("    ]\n")
            mfile.write("}\n")
            mfile.write("}\n")



def manuscript_to_csv( mdir, mfile, afile, ofile ):
    global manuscript

    set_manuscriptdir(mdir)
    set_manuscriptfile(mfile)
    set_authorfile(afile)
    read_data()

    names = ["title", "pov", "tod", "setting", "desc"]

    with open(ofile, "w") as ofile:
        first = True
        for name in names:
            if first:
                first = False
            else:
                ofile.write(",")
            ofile.write(name)

        ofile.write("\n")

        for chapter in manuscript["chapters"]:
            first = True
            for name in names:
                if first:
                    first = False
                else:
                    ofile.write(",")
                if name in chapter:  
                    value = chapter[name]
                else:
                    value = ""
                ofile.write("\"{}\"".format(value))

            ofile.write("\n")
            

def oms_to_docx(msdir, msfile, afile, outputfile):
    global author

    doc = Document()
    set_manuscriptdir(msdir)
    read_data()

    p = doc.add_paragraph("{}\n{}\n{}, {} {}\n{}\n{}\n".format(
                            author["name"],
                            author["streetAddress"],
                            author["addressLocality"], 
                            author["addressRegion"], 
                            author["postalCode"],
                            author["phone"],
                            author["email"]))

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.alignment = WD_LINE_SPACING.SINGLE

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph("{}\nby {}".format( manuscript["title"], 
                            author["name"]))
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.alignment = WD_LINE_SPACING.SINGLE

    doc.save(outputfile)


