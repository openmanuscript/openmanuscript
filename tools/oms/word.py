from . import core

import os
import json
import csv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

def kick_tires():
    print("oms.word test print")
    print("    author: {}".format(core.author))

def write_docx(msdir, msfile, afile, outputfile):
    doc = Document()
    core.set_manuscriptdir(msdir)
    core.read_data()

    p = doc.add_paragraph("{}\n{}\n{}, {} {}\n{}\n{}\n".format(
                            core.author["name"],
                            core.author["streetAddress"],
                            core.author["addressLocality"], 
                            core.author["addressRegion"], 
                            core.author["postalCode"],
                            core.author["phone"],
                            core.author["email"]))

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph("{}\nby {}".format( core.manuscript["title"], 
                            core.author["name"]))
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    doc.add_page_break()

    doc.add_paragraph("Test paragraph.")

    doc.save(outputfile)


